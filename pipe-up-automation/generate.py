#!/usr/bin/env python3
"""
Pipe-Up Daily Compliance Report Generator
==========================================
Takes a daily work plan PDF and a KML pipeline route file, and generates:
1. An interactive HTML map showing crew locations on the pipeline route
2. A Word document compliance report cross-referencing crew locations against regulatory zones

Usage:
    python generate.py                      # Process most recent PDF
    python generate.py --date 2026-03-03    # Process PDF for specific date
"""

import argparse
import json
import math
import os
import re
import sys
from datetime import datetime, timedelta
from pathlib import Path

import pdfplumber
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ─── Configuration ───────────────────────────────────────────────────────────

BASE_DIR = Path(__file__).parent
CONFIG = json.loads((BASE_DIR / "config.json").read_text())

KP_MIN, KP_MAX = CONFIG["kp_range"]

ACTIVITY_CATEGORIES = {
    "backfill":  {"label": "Backfill",         "color": "#2196F3"},
    "ditch":     {"label": "Ditching",         "color": "#FF9800"},
    "grad":      {"label": "Grading/Blasting", "color": "#795548"},
    "blast":     {"label": "Grading/Blasting", "color": "#795548"},
    "weld":      {"label": "Welding",          "color": "#F44336"},
    "stove":     {"label": "Welding",          "color": "#F44336"},
    "bend":      {"label": "Welding",          "color": "#F44336"},
    "coat":      {"label": "Coating",          "color": "#9C27B0"},
    "jeep":      {"label": "Coating",          "color": "#9C27B0"},
    "drill":     {"label": "Drilling",         "color": "#607D8B"},
    "split":     {"label": "Drilling",         "color": "#607D8B"},
    "grout":     {"label": "Drilling",         "color": "#607D8B"},
    "nail":      {"label": "Drilling",         "color": "#607D8B"},
    "crush":     {"label": "Crushing",         "color": "#FF5722"},
    "pump":      {"label": "Pump Support",     "color": "#00BCD4"},
    "tie":       {"label": "Tie-in",           "color": "#E91E63"},
    "ecb":       {"label": "Civil/Roads",      "color": "#3F51B5"},
    "road":      {"label": "Civil/Roads",      "color": "#3F51B5"},
    "barricade": {"label": "Civil/Roads",      "color": "#3F51B5"},
    "bypass":    {"label": "Civil/Roads",      "color": "#3F51B5"},
    "haul":      {"label": "Hauling",          "color": "#8BC34A"},
    "truck":     {"label": "Hauling",          "color": "#8BC34A"},
    "test":      {"label": "Testing/QC",       "color": "#FFC107"},
    "wqt":       {"label": "Testing/QC",       "color": "#FFC107"},
    "string":    {"label": "Stringing",        "color": "#26A69A"},
    "hydrovac":  {"label": "Hydrovac",         "color": "#7E57C2"},
    "desc":      {"label": "DESC/ESC",         "color": "#78909C"},
}

DEFAULT_CATEGORY = {"label": "Other", "color": "#9E9E9E"}

ZONE_TYPE_COLORS = {
    "fisheries":         "#00BCD4",
    "environmental":     "#FF9800",
    "ground_disturbance": "#4CAF50",
    "invasive_species":  "#E91E63",
    "safety":            "#F44336",
}

ZONE_TYPE_ICONS = {
    "fisheries":         "\U0001F41F",   # 🐟
    "environmental":     "\u26F0\uFE0F", # ⛰️
    "ground_disturbance": "\U0001F4CB",  # 📋
    "invasive_species":  "\u2623\uFE0F", # ☣️
    "safety":            "\U0001F4A5",   # 💥
}

HIGH_STATUSES = {"PENDING", "RESTRICTION ACTIVE", "ACTIVE TODAY", "CLOSED"}
MONITOR_STATUSES = {"ACTIVE WORK", "ACTIVE HAULING", "MONITORING"}
COMPLIANT_STATUSES = {"VALID", "OPEN"}

# Fixed section order matching reference document
SECTION_ORDER = ["fisheries", "environmental", "ground_disturbance", "safety", "invasive_species"]

SECTION_CONFIG = {
    "fisheries": {
        "heading": "FISHERIES TIMING WINDOWS (DFO)",
        "first_col": "CROSSING",
        "columns": ["CROSSING", "KP RANGE", "STATUS", "RESTRICTION", "AUTHORITY"],
        "format": "table",
        "has_crew_table": True,
    },
    "environmental": {
        "heading": "ENVIRONMENTAL SENSITIVE AREAS",
        "first_col": "ESA",
        "columns": ["ESA", "KP RANGE", "STATUS", "RESTRICTION", "AUTHORITY"],
        "format": "table",
        "has_crew_table": True,
    },
    "ground_disturbance": {
        "heading": "GROUND DISTURBANCE PERMITS",
        "first_col": "PERMIT",
        "columns": ["PERMIT", "KP RANGE", "STATUS", "CONDITIONS", "AUTHORITY"],
        "format": "table",
        "has_crew_table": False,
    },
    "safety": {
        "heading": "SAFETY EXCLUSION ZONES",
        "first_col": "ZONE",
        "columns": [],
        "format": "paragraph",
        "has_crew_table": False,
    },
    "invasive_species": {
        "heading": "INVASIVE SPECIES MANAGEMENT",
        "first_col": "ZONE",
        "columns": [],
        "format": "paragraph",
        "has_crew_table": False,
    },
}


# ─── Step 1: Find the Latest PDF ────────────────────────────────────────────

def find_pdf(target_date=None):
    """Find the PDF to process. If target_date given, match by filename date."""
    reports_dir = BASE_DIR / CONFIG["reports_dir"]
    pdfs = sorted(reports_dir.glob("*.pdf"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not pdfs:
        print("ERROR: No PDF files found in daily_reports/")
        sys.exit(1)

    if target_date:
        # Try to match date in filename
        for pdf in pdfs:
            if target_date.strftime("%B") in pdf.name and str(target_date.day) in pdf.name and str(target_date.year) in pdf.name:
                return pdf, target_date
            if target_date.strftime("%Y-%m-%d") in pdf.name:
                return pdf, target_date
        print(f"WARNING: No PDF found matching date {target_date.strftime('%Y-%m-%d')}. Using most recent.")

    pdf = pdfs[0]
    # Try to extract date from filename
    # Format: EGMP_Daily_Work_Plan_2330_-_Month_DD__YYYY.pdf
    # Also: EGMP Daily Work Plan 2330 - March 03, 2026.pdf
    date_match = re.search(r'(\w+)\s+(\d{1,2}),?\s+(\d{4})', pdf.name)
    if date_match:
        try:
            report_date = datetime.strptime(f"{date_match.group(1)} {date_match.group(2)} {date_match.group(3)}", "%B %d %Y").date()
            return pdf, report_date
        except ValueError:
            pass
    # Also try underscore format
    date_match = re.search(r'(\w+)_(\d{1,2})__(\d{4})', pdf.name)
    if date_match:
        try:
            report_date = datetime.strptime(f"{date_match.group(1)} {date_match.group(2)} {date_match.group(3)}", "%B %d %Y").date()
            return pdf, report_date
        except ValueError:
            pass

    return pdf, datetime.now().date()


# ─── Step 2: Parse the KML File ─────────────────────────────────────────────

def parse_kml():
    """Parse the KML to extract chainage points and centerline coordinates.
    Caches result as kml_cache.json for speed on repeat runs."""
    cache_path = BASE_DIR / "data" / "kml_cache.json"
    kml_path = BASE_DIR / CONFIG["kml_path"]

    if cache_path.exists() and cache_path.stat().st_mtime > kml_path.stat().st_mtime:
        data = json.loads(cache_path.read_text())
        return {float(k): v for k, v in data["kp_points"].items()}, data["centerline"]

    print("  Parsing KML (first run — will cache)...")
    content = kml_path.read_text()

    # Extract chainage placemarks (welds, bends, bore faces, open ends)
    # Names like "32+674.54", "0+012.76"
    kp_points = {}
    placemark_pattern = re.compile(
        r'<Placemark>\s*<name>(\d+\+[\d.]+)</name>.*?<Point>\s*<coordinates>\s*([-\d.]+),([-\d.]+)',
        re.DOTALL
    )
    for m in placemark_pattern.finditer(content):
        chainage_str = m.group(1)
        lng, lat = float(m.group(2)), float(m.group(3))
        parts = chainage_str.split('+')
        kp_val = int(parts[0]) + float(parts[1]) / 1000.0
        if KP_MIN <= kp_val <= KP_MAX + 0.5:
            kp_points[round(kp_val, 4)] = [lat, lng]

    # Extract centerline coordinates from all LineString elements inside Centerline folder
    cl_section_match = re.search(
        r'<Folder>\s*<name>Centerline</name>(.*?)</Folder>',
        content, re.DOTALL
    )
    centerline = []
    if cl_section_match:
        cl_section = cl_section_match.group(1)
        coord_blocks = re.findall(r'<coordinates>\s*(.*?)\s*</coordinates>', cl_section, re.DOTALL)
        for block in coord_blocks:
            for pair in block.strip().split():
                parts = pair.split(',')
                if len(parts) >= 2:
                    lng, lat = float(parts[0]), float(parts[1])
                    # Deduplicate consecutive identical points
                    if not centerline or (centerline[-1] != [lat, lng]):
                        centerline.append([lat, lng])

    print(f"    KP reference points: {len(kp_points)}")
    print(f"    Centerline vertices: {len(centerline)}")

    # Cache for next time
    cache_data = {"kp_points": {str(k): v for k, v in kp_points.items()}, "centerline": centerline}
    cache_path.write_text(json.dumps(cache_data))

    return {float(k): v for k, v in kp_points.items()}, centerline


# ─── Step 3 & 4: Parse the PDF and Extract KP Locations ─────────────────────

def parse_pdf(pdf_path):
    """Parse the daily work plan PDF and extract crew data with KP locations."""
    crews = []

    with pdfplumber.open(str(pdf_path)) as pdf:
        all_rows = []
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                all_rows.extend(table)

    current_inspector = None
    current_crew = None
    current_activity_parts = []

    def flush_crew():
        """Save the accumulated crew data."""
        nonlocal current_inspector, current_crew, current_activity_parts
        if current_inspector and current_crew and current_activity_parts:
            activity_text = "\n".join(current_activity_parts).strip()
            kp_locations = extract_kp_locations(activity_text)
            category = categorize_activity(activity_text, current_crew)
            crews.append({
                "inspector": current_inspector,
                "crew": current_crew,
                "activity": activity_text,
                "kp_locations": kp_locations,
                "category": category,
            })

    for row in all_rows:
        if not row or len(row) < 4:
            continue

        col0 = (row[0] or "").strip()
        col1 = (row[1] or "").strip()
        col3 = (row[3] or "").strip()

        # Skip header/title rows
        if not col0 and not col3:
            continue
        if "DISCIPLINE" in col1:
            continue
        if "SMJV" in col0 or "Eagle Mountain" in col0 or "Daily Activity" in col0:
            continue

        # New crew entry (has inspector name in col0)
        if col0 and col1 and not any(skip in col0 for skip in ["SMJV", "Eagle", "Daily", "Fortis"]):
            flush_crew()
            current_inspector = col0.replace("\n", " ").strip()
            # Extract crew name from col1 (first line, before Sub-Contractors/Employees)
            crew_lines = col1.split("\n")
            current_crew = crew_lines[0].strip()
            current_activity_parts = [col3] if col3 else []
        elif col3 and current_inspector:
            # Continuation row with more activity text
            current_activity_parts.append(col3)

    flush_crew()

    # Filter to only crews with at least one KP location
    crews_with_kp = [c for c in crews if c["kp_locations"]]
    crews_without_kp = [c for c in crews if not c["kp_locations"]]

    return crews_with_kp, crews_without_kp


def extract_kp_locations(text):
    """Extract KP values from activity text using multiple patterns."""
    kp_values = set()

    # Pattern 1: "KP 14+638", "kp 2+700", "Kp 29+600"
    for m in re.finditer(r'(?:KP|Kp|kp)\s*(\d{1,3})\+(\d{1,4})', text):
        kp = int(m.group(1)) + float(m.group(2)) / 1000.0
        if KP_MIN <= kp <= KP_MAX:
            kp_values.add(round(kp, 3))

    # Pattern 2: "5+222", "19+195" (no KP prefix)
    for m in re.finditer(r'(?<!\d)(\d{1,2})\+(\d{2,4})(?!\d)', text):
        kp = int(m.group(1)) + float(m.group(2)) / 1000.0
        if KP_MIN <= kp <= KP_MAX:
            kp_values.add(round(kp, 3))

    # Pattern 3: "KP 14 yard", "kp14", "KP 14" (whole KP, no chainage)
    for m in re.finditer(r'(?:KP|Kp|kp)\s*(\d{1,3})(?:\s|$|[^+\d])', text):
        kp = float(m.group(1))
        if KP_MIN <= kp <= KP_MAX:
            kp_values.add(round(kp, 3))

    # Pattern 4: "km 4.2", "KM 3.6"
    for m in re.finditer(r'(?:km|KM|Km)\s*(\d{1,3}(?:\.\d+)?)', text):
        kp = float(m.group(1))
        if KP_MIN <= kp <= KP_MAX:
            kp_values.add(round(kp, 3))

    return sorted(kp_values)


# ─── Step 5: Interpolate KP Values to Coordinates ───────────────────────────

def build_kp_interpolator(kp_points):
    """Build an interpolation function from KP value to [lat, lng]."""
    sorted_kps = sorted(kp_points.keys())

    def interpolate(target_kp):
        if target_kp <= sorted_kps[0]:
            return kp_points[sorted_kps[0]]
        if target_kp >= sorted_kps[-1]:
            return kp_points[sorted_kps[-1]]

        # Binary search for bracketing KPs
        lo, hi = 0, len(sorted_kps) - 1
        while lo < hi - 1:
            mid = (lo + hi) // 2
            if sorted_kps[mid] <= target_kp:
                lo = mid
            else:
                hi = mid

        lower_kp = sorted_kps[lo]
        upper_kp = sorted_kps[hi]
        if upper_kp == lower_kp:
            return kp_points[lower_kp]

        t = (target_kp - lower_kp) / (upper_kp - lower_kp)
        lower_coord = kp_points[lower_kp]
        upper_coord = kp_points[upper_kp]
        return [
            lower_coord[0] + t * (upper_coord[0] - lower_coord[0]),
            lower_coord[1] + t * (upper_coord[1] - lower_coord[1]),
        ]

    return interpolate


# ─── Step 6: Categorize Activities ──────────────────────────────────────────

def categorize_activity(activity_text, crew_name=""):
    """Assign activity category based on keywords. Crew name takes priority."""
    crew_lower = crew_name.lower()
    activity_lower = activity_text.lower()
    # Check crew name first — crew name is the strongest signal
    for keyword, cat in ACTIVITY_CATEGORIES.items():
        if keyword in crew_lower:
            return cat
    # Fall back to activity text
    for keyword, cat in ACTIVITY_CATEGORIES.items():
        if keyword in activity_lower:
            return cat
    return DEFAULT_CATEGORY


# ─── Step 7: Cross-Reference Against Regulatory Zones ───────────────────────

def load_zones():
    """Load regulatory zones from JSON."""
    zones_path = BASE_DIR / CONFIG["zones_path"]
    data = json.loads(zones_path.read_text())
    return data["zones"], data.get("zone_type_config", {})


def cross_reference_zones(crews, zones, interpolate_kp):
    """Check each crew KP against regulatory zones. Return findings."""
    findings = []
    for crew in crews:
        for kp in crew["kp_locations"]:
            for zone in zones:
                if zone["kp_start"] <= kp <= zone["kp_end"]:
                    status = zone["status"].upper()
                    if status in HIGH_STATUSES or zone.get("status_detail"):
                        severity = "HIGH"
                    elif status in MONITOR_STATUSES:
                        severity = "MONITOR"
                    else:
                        severity = "COMPLIANT"

                    findings.append({
                        "crew": crew["crew"],
                        "inspector": crew["inspector"],
                        "activity": crew["activity"],
                        "kp": kp,
                        "coord": interpolate_kp(kp),
                        "zone_name": zone["name"],
                        "zone_type": zone["type"],
                        "zone_status": zone["status"],
                        "zone_restriction": zone["restriction"],
                        "zone_authority": zone["authority"],
                        "severity": severity,
                    })
    return findings


# ─── Step 8: Generate Compliance Alerts ──────────────────────────────────────

def _find_crews_in_zone(crews, zone):
    """Find crews with KP locations inside a zone's KP range."""
    matches = []
    for crew in crews:
        for kp in crew["kp_locations"]:
            if zone["kp_start"] <= kp <= zone["kp_end"]:
                matches.append({"crew": crew["crew"], "kp": kp,
                                "inspector": crew.get("inspector", ""),
                                "activity": crew.get("category", {}).get("label", "")})
                break  # one match per crew is enough
    return matches


def generate_alerts(crews, zones, report_date, interpolate_kp):
    """Generate compliance alerts based on zone/crew intersections."""
    alerts = []

    # Build set of KPs where crews are active
    crew_kps = set()
    for crew in crews:
        crew_kps.update(crew["kp_locations"])

    for zone in zones:
        zone_crew_kps = [kp for kp in crew_kps if zone["kp_start"] <= kp <= zone["kp_end"]]
        if not zone_crew_kps:
            continue

        status = zone["status"].upper()
        zone_crews = _find_crews_in_zone(crews, zone)
        crew_snippet = ""
        if zone_crews:
            c = zone_crews[0]
            crew_snippet = f"{c['activity']} crew ({c['inspector']} / {c['crew']}) working at KP {c['kp']:.3f}. "

        # Fisheries zone with status_detail (closing soon)
        if zone["type"] == "fisheries" and zone.get("status_detail"):
            days_match = re.search(r'(\d+)\s*DAY', zone.get("status_detail", ""), re.IGNORECASE)
            days_left = int(days_match.group(1)) if days_match else None
            if days_left is not None and days_left <= 7:
                detail = (f"{crew_snippet}"
                          f"Verify all in-stream work completes before deadline. "
                          f"{zone['restriction']}.")
                alerts.append({
                    "severity": "HIGH",
                    "title": f"Fisheries Window Closing \u2014 {days_left} Days",
                    "detail": detail,
                    "kp": zone["kp_start"],
                    "coord": interpolate_kp(zone["kp_start"]),
                })

        # Pending GD zone with crew work
        if zone["type"] == "ground_disturbance" and status == "PENDING":
            detail = (f"{crew_snippet}"
                      f"GD permit pending \u2014 crews operating in zone without final sign-off. "
                      f"{zone['restriction']}.")
            alerts.append({
                "severity": "MEDIUM",
                "title": "GD Pending \u2014 Active Work in Zone",
                "detail": detail,
                "kp": zone["kp_start"],
                "coord": interpolate_kp(zone["kp_start"]),
            })

        # Archaeological monitor for GD zones with arch conditions
        if zone["type"] == "ground_disturbance" and status == "VALID":
            restriction_lower = zone.get("restriction", "").lower()
            if "archaeolog" in restriction_lower:
                detail = (f"{crew_snippet}"
                          f"Archaeological monitoring required per permit conditions. "
                          f"{zone['restriction']}.")
                alerts.append({
                    "severity": "MEDIUM",
                    "title": "Archaeological Monitor Zone",
                    "detail": detail,
                    "kp": zone["kp_start"],
                    "coord": interpolate_kp(zone["kp_start"]),
                })

        # Active safety zone
        if zone["type"] == "safety" and "ACTIVE" in status:
            detail = (f"{crew_snippet}"
                      f"{zone['restriction']}.")
            alerts.append({
                "severity": "MEDIUM",
                "title": f"Active Safety Zone \u2014 {zone['name']}",
                "detail": detail,
                "kp": zone["kp_start"],
                "coord": interpolate_kp(zone["kp_start"]),
            })

        # Environmental restriction in effect with activity
        if zone["type"] == "environmental" and status in {"RESTRICTION ACTIVE", "ACTIVE WORK"}:
            detail = (f"{crew_snippet}"
                      f"{zone['restriction']}.")
            alerts.append({
                "severity": "LOW",
                "title": f"Environmental Restriction \u2014 {zone['name']}",
                "detail": detail,
                "kp": zone["kp_start"],
                "coord": interpolate_kp(zone["kp_start"]),
            })

    # Sort by severity
    severity_order = {"HIGH": 0, "MEDIUM": 1, "LOW": 2}
    alerts.sort(key=lambda a: severity_order.get(a["severity"], 3))
    return alerts


# ─── Step 9: Generate the HTML Map ──────────────────────────────────────────

def generate_html_map(crews, zones, zone_config, findings, alerts, centerline,
                      kp_points, interpolate_kp, report_date, pdf_name, output_path):
    """Generate a self-contained interactive HTML map with Leaflet.js.
    Matches the visual style of EGP_Regulatory_Compliance_Map.html and EGP_Daily_Map_Feb27.html."""

    # Build crew markers data
    crew_markers = []
    for crew in crews:
        for kp in crew["kp_locations"]:
            coord = interpolate_kp(kp)
            crew_markers.append({
                "lat": coord[0], "lng": coord[1], "kp": kp,
                "inspector": crew["inspector"], "crew": crew["crew"],
                "activity": crew["activity"][:200], "color": crew["category"]["color"],
                "category": crew["category"]["label"],
            })

    # Build zone segments with coords, icon, and proper colors
    zone_segments = []
    sorted_kps = sorted(kp_points.keys())
    for zone in zones:
        zone_coords = []
        start_coord = interpolate_kp(zone["kp_start"])
        end_coord = interpolate_kp(zone["kp_end"])
        zone_coords.append(start_coord)
        for kp_val in sorted_kps:
            if zone["kp_start"] < kp_val < zone["kp_end"]:
                zone_coords.append(kp_points[kp_val])
        zone_coords.append(end_coord)

        # PENDING ground disturbance uses yellow color
        if zone["status"].upper() == "PENDING" and zone["type"] == "ground_disturbance":
            color = "#FFC107"
            icon = "\u26A0\uFE0F"  # ⚠️
        else:
            color = ZONE_TYPE_COLORS.get(zone["type"], "#9E9E9E")
            icon = ZONE_TYPE_ICONS.get(zone["type"], "")

        status = zone["status"].upper()
        is_warning = status in HIGH_STATUSES or status in MONITOR_STATUSES

        zone_segments.append({
            "name": zone["name"], "type": zone["type"],
            "kp_start": zone["kp_start"], "kp_end": zone["kp_end"],
            "status": zone["status"],
            "status_detail": zone.get("status_detail", ""),
            "restriction": zone["restriction"],
            "authority": zone["authority"],
            "color": color, "icon": icon,
            "is_warning": is_warning,
            "coords": zone_coords,
            "label": zone_config.get(zone["type"], {}).get("label", zone["type"]),
        })

    # KP whole-number markers
    kp_markers_data = []
    for kp_int in range(int(KP_MIN), int(KP_MAX) + 1):
        coord = interpolate_kp(float(kp_int))
        kp_markers_data.append({"kp": kp_int, "lat": coord[0], "lng": coord[1]})

    # Stats
    total_crews = len(crews)
    total_kps = sum(len(c["kp_locations"]) for c in crews)
    high_count = sum(1 for a in alerts if a["severity"] == "HIGH")
    med_count = sum(1 for a in alerts if a["severity"] == "MEDIUM")
    low_count = sum(1 for a in alerts if a["severity"] == "LOW")

    # Build unique activity legend
    legend_items = {}
    for crew in crews:
        cat = crew["category"]
        legend_items[cat["label"]] = cat["color"]

    # Zone type config for JS (with icons)
    zone_type_js = {}
    for ztype, zconf in zone_config.items():
        zone_type_js[ztype] = {
            "label": zconf.get("label", ztype),
            "color": ZONE_TYPE_COLORS.get(ztype, "#9E9E9E"),
            "icon": ZONE_TYPE_ICONS.get(ztype, ""),
        }

    date_str = report_date.strftime("%B %d, %Y")

    # Alert badges HTML
    badge_html = ""
    if high_count:
        badge_html += f'<div class="alert-badge high">\u26A0 {high_count} HIGH</div>'
    if med_count:
        badge_html += f'<div class="alert-badge medium">\u26A0 {med_count} MEDIUM</div>'
    if low_count:
        badge_html += f'<div class="alert-badge low">\u2139 {low_count} INFO</div>'

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>EGP Regulatory Compliance Map \u2014 {date_str}</title>
<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"/>
<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
<style>
* {{ box-sizing: border-box; margin: 0; padding: 0; }}
body {{ font-family: 'Segoe UI', system-ui, sans-serif; background: #0f1923; color: #e0e0e0; }}

.header {{
  background: linear-gradient(135deg, #0B1D33 0%, #132B4A 100%);
  padding: 12px 24px; display: flex; justify-content: space-between; align-items: center;
  border-bottom: 2px solid #E8913A;
}}
.header h1 {{ font-size: 1.05rem; color: #fff; font-weight: 700; line-height: 1.4; }}
.header h1 span {{ color: #E8913A; }}
.alert-summary {{ display: flex; gap: 12px; align-items: center; }}
.alert-badge {{
  display: flex; align-items: center; gap: 6px; padding: 6px 14px;
  border-radius: 20px; font-size: 0.75rem; font-weight: 700;
}}
.alert-badge.high {{ background: rgba(244,67,54,0.15); color: #E57373; border: 1px solid rgba(244,67,54,0.3); }}
.alert-badge.medium {{ background: rgba(255,152,0,0.15); color: #FFB74D; border: 1px solid rgba(255,152,0,0.3); }}
.alert-badge.low {{ background: rgba(33,150,243,0.15); color: #64B5F6; border: 1px solid rgba(33,150,243,0.3); }}

.main {{ display: flex; height: calc(100vh - 56px); }}
#map {{ flex: 1; z-index: 1; }}

.sidebar {{
  width: 380px; background: #0f1923; overflow-y: auto;
  border-left: 1px solid #1e2d3d; padding: 0;
}}
.sidebar-section {{ padding: 14px 16px; border-bottom: 1px solid #1e2d3d; }}
.sidebar h2 {{
  font-size: 0.72rem; text-transform: uppercase; letter-spacing: 1.5px;
  color: #E8913A; margin-bottom: 10px;
}}

/* Zone toggles */
.zone-toggle {{
  display: flex; align-items: center; gap: 10px; padding: 7px 0;
  cursor: pointer; font-size: 0.8rem; transition: opacity 0.2s;
}}
.zone-toggle.off {{ opacity: 0.3; }}
.zone-bar {{ width: 24px; height: 6px; border-radius: 3px; flex-shrink: 0; }}
.zone-count {{ margin-left: auto; color: #5a6978; font-size: 0.7rem; }}

/* Alert cards */
.alert-card {{
  background: #162233; border: 1px solid #1e2d3d; border-radius: 8px;
  padding: 12px; margin-bottom: 8px; cursor: pointer; transition: all 0.2s;
}}
.alert-card:hover {{ background: #1a2a3e; }}
.alert-card.severity-HIGH {{ border-left: 3px solid #F44336; }}
.alert-card.severity-MEDIUM {{ border-left: 3px solid #FF9800; }}
.alert-card.severity-LOW {{ border-left: 3px solid #2196F3; }}
.alert-title {{ font-size: 0.82rem; font-weight: 700; color: #fff; margin-bottom: 4px; display: flex; align-items: center; gap: 6px; }}
.alert-detail {{ font-size: 0.72rem; color: #8a9bb0; line-height: 1.4; }}
.alert-kp {{ font-size: 0.7rem; color: #E8913A; margin-top: 6px; font-weight: 600; }}
.alert-sev {{
  font-size: 0.6rem; font-weight: 700; padding: 2px 6px; border-radius: 3px;
  display: inline-block; margin-top: 4px;
}}
.alert-sev.HIGH {{ background: rgba(244,67,54,0.15); color: #E57373; }}
.alert-sev.MEDIUM {{ background: rgba(255,152,0,0.15); color: #FFB74D; }}
.alert-sev.LOW {{ background: rgba(33,150,243,0.15); color: #64B5F6; }}

/* Zone detail cards */
.zone-card {{
  background: #162233; border: 1px solid #1e2d3d; border-radius: 8px;
  padding: 10px 12px; margin-bottom: 6px; cursor: pointer; transition: all 0.2s;
  border-left: 3px solid transparent;
}}
.zone-card:hover {{ background: #1a2a3e; }}
.zc-name {{ font-size: 0.8rem; font-weight: 600; color: #fff; }}
.zc-range {{ font-size: 0.7rem; color: #E8913A; margin-top: 2px; }}
.zc-restriction {{ font-size: 0.7rem; color: #8a9bb0; margin-top: 4px; line-height: 1.35; }}
.zc-authority {{ font-size: 0.65rem; color: #5a6978; margin-top: 4px; font-style: italic; }}
.zc-status {{
  display: inline-block; font-size: 0.6rem; font-weight: 700; padding: 2px 6px;
  border-radius: 3px; margin-top: 4px; text-transform: uppercase;
}}
.zc-status.VALID, .zc-status.OPEN, .zc-status.MONITORING {{ background: rgba(76,175,80,0.15); color: #81C784; }}
.zc-status.ACTIVE {{ background: rgba(255,152,0,0.15); color: #FFB74D; }}
.zc-status.PENDING {{ background: rgba(255,193,7,0.15); color: #FFD54F; }}
.zc-status.RESTRICTED, .zc-status.CLOSED {{ background: rgba(244,67,54,0.15); color: #E57373; }}

/* Crew toggle */
.crew-toggle {{
  display: flex; align-items: center; gap: 8px; padding: 8px 12px;
  background: #162233; border: 1px solid #1e2d3d; border-radius: 8px;
  cursor: pointer; font-size: 0.8rem; margin-bottom: 8px; transition: all 0.2s;
}}
.crew-toggle.active {{ border-color: #E8913A; }}
.crew-dot {{ width: 10px; height: 10px; border-radius: 50%; background: #E8913A; }}

/* Activity legend */
.legend-item {{
  display: flex; align-items: center; gap: 8px;
  padding: 6px 0; font-size: 0.8rem; cursor: pointer;
  opacity: 1; transition: opacity 0.2s;
}}
.legend-item.dimmed {{ opacity: 0.3; }}
.legend-dot {{ width: 12px; height: 12px; border-radius: 50%; flex-shrink: 0; }}

/* Leaflet popups */
.leaflet-popup-content-wrapper {{ background: #162233; color: #e0e0e0; border-radius: 8px; max-width: 320px; }}
.leaflet-popup-tip {{ background: #162233; }}
.popup-zone-title {{ font-weight: 700; font-size: 0.9rem; margin-bottom: 6px; }}
.popup-zone-type {{ font-size: 0.7rem; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 6px; font-weight: 600; }}
.popup-zone-restriction {{ font-size: 0.78rem; color: #8a9bb0; line-height: 1.4; margin-bottom: 6px; }}
.popup-zone-authority {{ font-size: 0.7rem; color: #5a6978; font-style: italic; }}
.popup-zone-status {{ display: inline-block; font-size: 0.7rem; font-weight: 700; padding: 2px 8px; border-radius: 4px; margin-top: 4px; }}

@keyframes pulse {{
  0%,100% {{ transform: scale(1); opacity: 1; }}
  50% {{ transform: scale(1.3); opacity: 0.7; }}
}}

@media (max-width: 768px) {{
  .sidebar {{ display: none; }}
  .main {{ flex-direction: column; }}
}}
</style>
</head>
<body>

<div class="header">
  <h1>EGP <span>Regulatory Compliance Map</span><br>
  <span style="font-size:0.78rem;color:#8a9bb0;font-weight:400;">{date_str} \u2014 {CONFIG['project_name']}</span></h1>
  <div class="alert-summary">{badge_html}</div>
</div>

<div class="main">
  <div id="map"></div>
  <div class="sidebar">
    <div class="sidebar-section">
      <h2>Compliance Alerts</h2>
      <div id="alertList"></div>
    </div>
    <div class="sidebar-section">
      <h2>Regulatory Layers</h2>
      <div id="zoneToggles"></div>
    </div>
    <div class="sidebar-section">
      <h2>Activity Legend</h2>
      <div id="activityLegend"></div>
    </div>
    <div class="sidebar-section">
      <h2>Crew Activity</h2>
      <div class="crew-toggle active" id="crewToggle" onclick="toggleCrews()">
        <div class="crew-dot"></div>
        <span>Show crew locations ({total_kps} markers)</span>
      </div>
    </div>
    <div class="sidebar-section">
      <h2>Zone Details</h2>
      <div id="zoneDetails"></div>
    </div>
  </div>
</div>

<script>
const zones = {json.dumps(zone_segments)};
const crewMarkers = {json.dumps(crew_markers)};
const alerts = {json.dumps(alerts)};
const centerline = {json.dumps(centerline)};
const kpTicks = {json.dumps(kp_markers_data)};
const legendItems = {json.dumps(legend_items)};
const zoneTypes = {json.dumps(zone_type_js)};

// Map
const map = L.map('map', {{ zoomControl: true, attributionControl: false }}).setView({CONFIG['map_center']}, {CONFIG['map_zoom']});
L.tileLayer('https://{{s}}.basemaps.cartocdn.com/dark_all/{{z}}/{{x}}/{{y}}{{r}}.png', {{ maxZoom: 19, subdomains: 'abcd' }}).addTo(map);

// Pipeline centerline
const pipeline = L.polyline(centerline, {{ color: '#3B8FD4', weight: 3, opacity: 0.4 }}).addTo(map);
map.fitBounds(pipeline.getBounds().pad(0.05));

// KP markers — divIcon labels
kpTicks.forEach(kp => {{
  L.marker([kp.lat, kp.lng], {{
    icon: L.divIcon({{
      className: '',
      html: `<div style="background:#1e2d3d;color:#6a7a8d;font-size:8px;padding:1px 3px;border-radius:2px;border:1px solid #2a3d52;white-space:nowrap;">KP ${{Math.round(kp.kp)}}</div>`,
      iconSize: null, iconAnchor: [12, 7]
    }})
  }}).addTo(map);
}});

// Draw regulatory zones
const zoneLayers = {{}};
const zoneVisible = {{}};

Object.keys(zoneTypes).forEach(type => {{
  zoneLayers[type] = [];
  zoneVisible[type] = true;
}});

zones.forEach(zone => {{
  if (!zone.coords || zone.coords.length < 2) return;
  const cfg = zoneTypes[zone.type] || {{ label: zone.type, color: '#999', icon: '' }};
  const isWarning = zone.is_warning;

  // Thick semi-transparent highlight
  const polyline = L.polyline(zone.coords, {{
    color: zone.color,
    weight: 14,
    opacity: 0.25,
    lineCap: 'round'
  }}).addTo(map);

  // Bright border
  const border = L.polyline(zone.coords, {{
    color: zone.color,
    weight: 4,
    opacity: isWarning ? 0.9 : 0.6,
    dashArray: isWarning ? '8, 6' : null
  }}).addTo(map);

  // Status color for popup badge
  let statusColor = '#81C784';
  let statusBg = 'rgba(76,175,80,0.2)';
  if (['PENDING', 'RESTRICTION ACTIVE', 'ACTIVE TODAY', 'ACTIVE WORK', 'ACTIVE HAULING'].includes(zone.status)) {{
    statusColor = '#FFB74D'; statusBg = 'rgba(255,152,0,0.2)';
  }}
  if (['CLOSED'].includes(zone.status)) {{
    statusColor = '#E57373'; statusBg = 'rgba(244,67,54,0.2)';
  }}

  const popup = `
    <div class="popup-zone-type" style="color:${{zone.color}}">${{zone.icon}} ${{zone.type.replace('_',' ').toUpperCase()}}</div>
    <div class="popup-zone-title" style="color:${{zone.color}}">${{zone.name}}</div>
    <div style="font-size:0.75rem;color:#E8913A;margin-bottom:6px;">KP ${{zone.kp_start.toFixed(1)}} \\u2014 KP ${{zone.kp_end.toFixed(1)}}</div>
    <div class="popup-zone-restriction">${{zone.restriction}}</div>
    <div class="popup-zone-authority">${{zone.authority}}</div>
    <div class="popup-zone-status" style="color:${{statusColor}};background:${{statusBg}}">${{zone.status}}${{zone.status_detail ? ' \\u2014 ' + zone.status_detail : ''}}</div>
  `;

  polyline.bindPopup(popup, {{ maxWidth: 300 }});
  border.bindPopup(popup, {{ maxWidth: 300 }});

  // Start marker — divIcon with colored bg, icon, KP label, glow
  const startIcon = L.divIcon({{
    className: '',
    html: `<div style="background:${{zone.color}};color:#fff;font-size:10px;padding:2px 6px;border-radius:4px;white-space:nowrap;font-weight:700;box-shadow:0 0 8px ${{zone.color}}44;">${{zone.icon}} KP ${{zone.kp_start.toFixed(1)}}</div>`,
    iconSize: null, iconAnchor: [0, 10]
  }});
  const startMarker = L.marker(zone.coords[0], {{ icon: startIcon }}).addTo(map).bindPopup(popup, {{ maxWidth: 300 }});

  if (zoneLayers[zone.type]) {{
    zoneLayers[zone.type].push(polyline, border, startMarker);
  }}
}});

// Crew markers — divIcon circles
const crewLayerGroup = [];
let crewsVisible = true;

crewMarkers.forEach(m => {{
  const icon = L.divIcon({{
    className: '',
    html: `<div style="width:10px;height:10px;background:${{m.color}};border:1.5px solid #fff;border-radius:50%;box-shadow:0 0 4px ${{m.color}}66;"></div>`,
    iconSize: [10, 10], iconAnchor: [5, 5]
  }});
  const popup = `<div style="font-weight:700;color:#E8913A;margin-bottom:4px;">${{m.inspector}}</div><div style="font-size:0.78rem;color:#8a9bb0;">${{m.crew}}</div><div style="font-size:0.78rem;margin-top:4px;">${{m.activity}}</div><div style="font-size:0.75rem;color:#E8913A;margin-top:4px;font-weight:600;">KP ${{m.kp.toFixed(3)}}</div><div style="font-size:0.7rem;color:#6a7a8d;margin-top:2px;">${{m.category}}</div>`;
  const marker = L.marker([m.lat, m.lng], {{ icon }}).addTo(map).bindPopup(popup);
  crewLayerGroup.push(marker);
}});

// Alert markers — divIcon with emoji + pulse
alerts.forEach(a => {{
  if (!a.coord) return;
  const pulse = a.severity === 'HIGH' ? 'animation:pulse 1.5s infinite;' : '';
  const color = a.severity === 'HIGH' ? '#F44336' : a.severity === 'MEDIUM' ? '#FF9800' : '#2196F3';
  const alertIcon = a.severity === 'HIGH' ? '\\u26a0\\ufe0f' : a.severity === 'MEDIUM' ? '\\u26a0\\ufe0f' : '\\u2139\\ufe0f';
  const icon = L.divIcon({{
    className: '',
    html: `<div style="width:24px;height:24px;display:flex;align-items:center;justify-content:center;background:${{color}}22;border:2px solid ${{color}};border-radius:50%;font-size:13px;${{pulse}}">${{alertIcon}}</div>`,
    iconSize: [24, 24], iconAnchor: [12, 12]
  }});
  const popup = `<div style="font-weight:700;color:${{color}};font-size:0.9rem;margin-bottom:4px;">${{alertIcon}} ${{a.title}}</div><div style="font-size:0.78rem;color:#8a9bb0;line-height:1.4;">${{a.detail}}</div><div style="font-size:0.75rem;color:#E8913A;margin-top:6px;font-weight:600;">KP ${{a.kp ? a.kp.toFixed(1) : ''}}</div>`;
  L.marker([a.coord[0], a.coord[1]], {{ icon, zIndexOffset: 1000 }}).addTo(map).bindPopup(popup, {{ maxWidth: 300 }});
}});

// ── Sidebar: Alert cards ──
const alertList = document.getElementById('alertList');
alerts.sort((a,b) => {{ const ord = {{HIGH:0,MEDIUM:1,LOW:2}}; return ord[a.severity]-ord[b.severity]; }});
alerts.forEach(a => {{
  const card = document.createElement('div');
  card.className = `alert-card severity-${{a.severity}}`;
  card.innerHTML = `
    <div class="alert-title">${{a.severity === 'HIGH' ? '\\u26a0\\ufe0f' : a.severity === 'MEDIUM' ? '\\u26a0\\ufe0f' : '\\u2139\\ufe0f'}} ${{a.title}}</div>
    <div class="alert-detail">${{a.detail}}</div>
    <div class="alert-kp">KP ${{a.kp ? a.kp.toFixed(1) : ''}}</div>
    <span class="alert-sev ${{a.severity}}">${{a.severity}}</span>
  `;
  card.onclick = () => {{
    if (a.coord) map.setView([a.coord[0], a.coord[1]], 14);
  }};
  alertList.appendChild(card);
}});

// ── Sidebar: Zone toggles ──
const togglesDiv = document.getElementById('zoneToggles');
Object.entries(zoneTypes).forEach(([type, cfg]) => {{
  const count = zones.filter(z => z.type === type).length;
  const toggle = document.createElement('div');
  toggle.className = 'zone-toggle';
  toggle.dataset.type = type;
  toggle.innerHTML = `<div class="zone-bar" style="background:${{cfg.color}}"></div>${{cfg.icon}} ${{cfg.label}}<span class="zone-count">${{count}}</span>`;
  toggle.onclick = () => {{
    zoneVisible[type] = !zoneVisible[type];
    toggle.classList.toggle('off');
    (zoneLayers[type] || []).forEach(layer => {{
      if (zoneVisible[type]) {{
        if (layer.setStyle) layer.setStyle({{ opacity: layer.options._origOpacity || 0.6 }});
        if (layer.getElement) {{ const el = layer.getElement(); if (el) el.style.display = ''; }}
        if (layer.setOpacity) layer.setOpacity(1);
      }} else {{
        if (layer.setStyle) layer.setStyle({{ opacity: 0 }});
        if (layer.getElement) {{ const el = layer.getElement(); if (el) el.style.display = 'none'; }}
        if (layer.setOpacity) layer.setOpacity(0);
      }}
    }});
  }};
  togglesDiv.appendChild(toggle);
}});

// ── Sidebar: Activity legend ──
const legendDiv = document.getElementById('activityLegend');
Object.entries(legendItems).sort().forEach(([label, color]) => {{
  const item = document.createElement('div');
  item.className = 'legend-item';
  item.innerHTML = `<div class="legend-dot" style="background:${{color}}"></div>${{label}}`;
  legendDiv.appendChild(item);
}});

// ── Crew toggle ──
function toggleCrews() {{
  crewsVisible = !crewsVisible;
  document.getElementById('crewToggle').classList.toggle('active');
  crewLayerGroup.forEach(m => {{
    if (crewsVisible) {{ m.setOpacity(1); if(m.getElement()) m.getElement().style.pointerEvents='auto'; }}
    else {{ m.setOpacity(0); if(m.getElement()) m.getElement().style.pointerEvents='none'; }}
  }});
}}

// ── Sidebar: Zone detail cards ──
const detailsDiv = document.getElementById('zoneDetails');
zones.forEach(zone => {{
  const cfg = zoneTypes[zone.type] || {{ label: zone.type, color: '#999', icon: '' }};
  let statusClass = 'VALID';
  if (['PENDING'].includes(zone.status)) statusClass = 'PENDING';
  if (['RESTRICTION ACTIVE', 'ACTIVE TODAY', 'ACTIVE WORK', 'ACTIVE HAULING'].includes(zone.status)) statusClass = 'ACTIVE';
  if (['OPEN', 'MONITORING'].includes(zone.status)) statusClass = zone.status;

  const card = document.createElement('div');
  card.className = 'zone-card';
  card.style.borderLeftColor = zone.color;
  card.innerHTML = `
    <div class="zc-name">${{zone.icon}} ${{zone.name}}</div>
    <div class="zc-range">KP ${{zone.kp_start.toFixed(1)}} \\u2014 KP ${{zone.kp_end.toFixed(1)}}</div>
    <div class="zc-restriction">${{zone.restriction}}</div>
    <div class="zc-authority">${{zone.authority}}</div>
    <span class="zc-status ${{statusClass}}">${{zone.status}}${{zone.status_detail ? ' \\u2014 ' + zone.status_detail : ''}}</span>
  `;
  card.onclick = () => {{
    if (zone.coords && zone.coords.length) {{
      const bounds = L.latLngBounds(zone.coords);
      map.fitBounds(bounds.pad(0.5));
    }}
  }};
  detailsDiv.appendChild(card);
}});
</script>
</body>
</html>"""

    output_path.write_text(html)


# ─── Step 10: Generate the Word Compliance Report ────────────────────────────

def generate_word_report(crews, zones, zone_config, findings, alerts,
                         report_date, pdf_name, output_path):
    """Generate a Word document compliance report matching reference format."""
    doc = Document()

    # ── Set default document font to Arial 11pt ──
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Arial"

    # ── Page setup: US Letter, portrait, 1" sides, 0.75" top/bottom ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    date_str = report_date.strftime("%B %d, %Y")

    # ── Header ──
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(f"CONFIDENTIAL \u2014 EGP Regulatory Compliance Report \u2014 {date_str}")
    run.font.size = Pt(8)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)

    # ── Footer with page number ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Pipe-Up Regulatory Compliance Engine | Page ")
    run.font.size = Pt(8)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    fld_instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run2 = fp.add_run()
    run2.font.size = Pt(8)
    run2.font.name = "Arial"
    run2.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)
    run2._r.append(fld_char_begin)
    run3 = fp.add_run()
    run3._r.append(fld_instr)
    run4 = fp.add_run()
    run4._r.append(fld_char_end)

    # ═══════════════════════════════════════════════
    # Page 1: Summary
    # ═══════════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(4)
    run = title.add_run("DAILY REGULATORY COMPLIANCE REPORT")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x0B, 0x1D, 0x33)

    # Subtitle line 1: project name
    sub1 = doc.add_paragraph()
    sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub1.space_after = Pt(2)
    run = sub1.add_run(CONFIG['project_name'])
    run.font.size = Pt(11)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x5A, 0x69, 0x78)

    # Subtitle line 2: date | job | location (pipe-separated)
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.space_after = Pt(6)
    run = sub2.add_run(
        f"{date_str}  |  {CONFIG['job_number']}  |  {CONFIG['location']}"
    )
    run.font.size = Pt(11)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x5A, 0x69, 0x78)

    doc.add_paragraph()  # spacer

    # Summary counts
    action_count = sum(1 for f in findings if f["severity"] == "HIGH")
    monitor_count = sum(1 for f in findings if f["severity"] == "MONITOR")
    compliant_count = sum(1 for f in findings if f["severity"] == "COMPLIANT")
    total_findings = len(findings)

    summary_table = doc.add_table(rows=2, cols=4)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(summary_table)
    _style_table_header(summary_table.rows[0],
                        ["TOTAL FINDINGS", "ACTION REQUIRED", "MONITOR", "COMPLIANT"])

    data_colors = [
        RGBColor(0x0B, 0x1D, 0x33),  # navy
        RGBColor(0xF4, 0x43, 0x36),  # red
        RGBColor(0xFF, 0x98, 0x00),  # amber
        RGBColor(0x4C, 0xAF, 0x50),  # green
    ]
    for i, val in enumerate([total_findings, action_count, monitor_count, compliant_count]):
        cell = summary_table.rows[1].cells[i]
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        run.font.size = Pt(20)
        run.font.name = "Arial"
        run.bold = True
        run.font.color.rgb = data_colors[i]

    doc.add_paragraph()

    # Compute crews operating in/adjacent to regulatory zones
    crews_in_zones = set()
    for f in findings:
        crews_in_zones.add(f["crew"])
    crews_in_zone_count = len(crews_in_zones)
    total_crew_count = len(crews)
    pct = round(100 * crews_in_zone_count / total_crew_count) if total_crew_count else 0

    summary_p = doc.add_paragraph()
    run = summary_p.add_run(
        f"Of {total_crew_count} active crews reporting today, "
        f"{crews_in_zone_count} crews ({pct}%) are operating within or adjacent to "
        f"one or more regulatory zones. A total of {total_findings} crew-zone intersections "
        f"were identified. {action_count} require immediate attention."
    )
    run.font.size = Pt(11)
    run.font.name = "Arial"

    # ═══════════════════════════════════════════════
    # Section 1: Compliance Alerts
    # ═══════════════════════════════════════════════
    doc.add_page_break()
    _add_section_heading(doc, "1. COMPLIANCE ALERTS \u2014 ACTION REQUIRED")

    if alerts:
        alert_table = doc.add_table(rows=1 + len(alerts), cols=4)
        _apply_table_borders(alert_table)
        _style_table_header(alert_table.rows[0], ["SEVERITY", "ALERT", "DETAIL", "LOCATION"])

        for row in alert_table.rows:
            row.cells[0].width = Cm(2.5)
            row.cells[1].width = Cm(4.5)
            row.cells[2].width = Cm(8.0)
            row.cells[3].width = Cm(2.0)

        for i, alert in enumerate(alerts):
            row = alert_table.rows[i + 1]
            row.cells[0].text = alert["severity"]
            row.cells[1].text = alert["title"]
            row.cells[2].text = alert["detail"][:200]
            row.cells[3].text = f"KP {alert.get('kp', 0):.1f}"

            sev_colors = {"HIGH": "F44336", "MEDIUM": "FF9800", "LOW": "2196F3"}
            _shade_cell(row.cells[0], sev_colors.get(alert["severity"], "999999"))
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            _set_row_font(row, Pt(9))
            if i % 2 == 1:
                _shade_row(row, "F5F5F5")
    else:
        doc.add_paragraph("No compliance alerts for this date.")

    # ═══════════════════════════════════════════════
    # Sections 2-6: Zone Sections (fixed order)
    # ═══════════════════════════════════════════════
    zone_types_present = set(z["type"] for z in zones)
    section_num = 2

    for ztype in SECTION_ORDER:
        if ztype not in zone_types_present:
            continue

        cfg = SECTION_CONFIG[ztype]
        type_zones = [z for z in zones if z["type"] == ztype]
        type_findings = [f for f in findings if f["zone_type"] == ztype]

        doc.add_page_break()
        _add_section_heading(doc, f"{section_num}. {cfg['heading']}")

        # Section intro paragraph
        intersection_count = len(type_findings)
        zone_count = len(type_zones)
        if ztype == "fisheries":
            intro_text = (f"{zone_count} fisheries timing windows are defined along the "
                          f"pipeline route. {intersection_count} active crew-zone "
                          f"intersections identified today.")
        elif ztype == "environmental":
            intro_text = (f"{zone_count} environmentally sensitive areas identified. "
                          f"{intersection_count} crew-zone intersections today.")
        elif ztype == "ground_disturbance":
            valid_count = sum(1 for z in type_zones if z["status"].upper() == "VALID")
            pending_count = sum(1 for z in type_zones if z["status"].upper() == "PENDING")
            intro_text = (f"{zone_count} ground disturbance permits tracked. "
                          f"{valid_count} valid, {pending_count} pending. "
                          f"{intersection_count} crew-zone intersections today.")
        elif ztype == "safety":
            active_count = sum(1 for z in type_zones if "ACTIVE" in z["status"].upper())
            intro_text = (f"{zone_count} safety exclusion zones defined. "
                          f"{active_count} currently active. "
                          f"{intersection_count} crew-zone intersections today.")
        elif ztype == "invasive_species":
            intro_text = (f"{zone_count} invasive species management zones along the "
                          f"pipeline route. {intersection_count} crew-zone intersections today.")
        else:
            intro_text = f"{zone_count} zones. {intersection_count} intersections today."

        ip = doc.add_paragraph()
        run = ip.add_run(intro_text)
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0x5A, 0x69, 0x78)
        run.italic = True
        ip.space_after = Pt(6)

        # ── Table format sections (fisheries, environmental, GD) ──
        if cfg["format"] == "table":
            zt = doc.add_table(rows=1 + len(type_zones), cols=5)
            _apply_table_borders(zt)
            _style_table_header(zt.rows[0], cfg["columns"])

            for i, z in enumerate(type_zones):
                row = zt.rows[i + 1]
                row.cells[0].text = z["name"]
                row.cells[1].text = f"KP {z['kp_start']:.1f} \u2014 {z['kp_end']:.1f}"

                # Status: single line with em dash
                status_text = z["status"]
                if z.get("status_detail"):
                    status_text += f" \u2014 {z['status_detail']}"
                row.cells[2].text = status_text

                row.cells[3].text = z["restriction"]
                row.cells[4].text = z["authority"]

                # Status cell conditional color
                status_upper = z["status"].upper()
                if status_upper in COMPLIANT_STATUSES:
                    _shade_cell(row.cells[2], "4CAF50")
                elif status_upper in HIGH_STATUSES:
                    _shade_cell(row.cells[2], "F44336")
                else:
                    _shade_cell(row.cells[2], "FF9800")
                row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                _set_row_font(row, Pt(9))
                if i % 2 == 1:
                    _shade_row(row, "F5F5F5")

            # Crew sub-table (only for fisheries and environmental)
            if cfg["has_crew_table"] and type_findings:
                type_label = zone_config.get(ztype, {}).get("label", ztype.replace("_", " ").title())
                doc.add_paragraph()
                p = doc.add_paragraph()
                run = p.add_run(f"Crews Operating in {type_label}:")
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = "Arial"
                run.font.color.rgb = RGBColor(0x0B, 0x1D, 0x33)

                ft = doc.add_table(rows=1 + len(type_findings), cols=4)
                _apply_table_borders(ft)
                _style_table_header(ft.rows[0], ["CREW", "KP", "ACTIVITY", "ZONE"])

                for i, f in enumerate(type_findings):
                    row = ft.rows[i + 1]
                    row.cells[0].text = f["crew"]
                    row.cells[1].text = f"KP {f['kp']:.3f}"
                    row.cells[2].text = f["activity"][:120]
                    row.cells[3].text = f["zone_name"]
                    _set_row_font(row, Pt(9))
                    if i % 2 == 1:
                        _shade_row(row, "F5F5F5")

        # ── Paragraph format sections (safety, invasive) ──
        elif cfg["format"] == "paragraph":
            for z in type_zones:
                zp = doc.add_paragraph()
                # Bold zone name
                run = zp.add_run(f"{z['name']} ")
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Arial"
                # KP range and details
                status_text = z["status"]
                if z.get("status_detail"):
                    status_text += f" \u2014 {z['status_detail']}"
                run = zp.add_run(
                    f"(KP {z['kp_start']:.1f} - {z['kp_end']:.1f}): "
                    f"{z['restriction']}. Status: {status_text}. "
                    f"Authority: {z['authority']}."
                )
                run.font.size = Pt(10)
                run.font.name = "Arial"
                zp.space_after = Pt(4)

            # Crew count note for invasive species
            if ztype == "invasive_species" and type_findings:
                crew_names = set(f["crew"] for f in type_findings)
                note_p = doc.add_paragraph()
                run = note_p.add_run(
                    f"{len(crew_names)} crew(s) operating in invasive species zones today. "
                    f"Standard wash-down and inspection protocols apply."
                )
                run.font.size = Pt(10)
                run.font.name = "Arial"
                run.italic = True

        section_num += 1

    # ═══════════════════════════════════════════════
    # Section 7: Complete Crew-Zone Intersection Log
    # ═══════════════════════════════════════════════
    doc.add_page_break()
    _add_section_heading(doc, f"{section_num}. COMPLETE CREW-ZONE INTERSECTION LOG")

    p = doc.add_paragraph()
    run = p.add_run(
        f"The following table lists every intersection between active crew locations and "
        f"regulatory zones for {date_str}. This log serves as the daily compliance record "
        f"for audit purposes."
    )
    run.font.size = Pt(9)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x5A, 0x69, 0x78)
    run.italic = True

    if findings:
        log_table = doc.add_table(rows=1 + len(findings), cols=5)
        _apply_table_borders(log_table)
        _style_table_header(log_table.rows[0],
                            ["CREW", "KP", "ZONE", "RESTRICTION", "STATUS"])

        for i, f in enumerate(findings):
            row = log_table.rows[i + 1]
            row.cells[0].text = f["crew"]
            row.cells[1].text = f"KP {f['kp']:.3f}"
            row.cells[2].text = f["zone_name"]
            row.cells[3].text = f["zone_restriction"]
            status_text = f["zone_status"]
            row.cells[4].text = status_text

            # Color the status cell
            status_upper = f["zone_status"].upper()
            if status_upper in COMPLIANT_STATUSES:
                _shade_cell(row.cells[4], "4CAF50")
            elif status_upper in HIGH_STATUSES:
                _shade_cell(row.cells[4], "F44336")
            else:
                _shade_cell(row.cells[4], "FF9800")
            row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            _set_row_font(row, Pt(8))
            if i % 2 == 1:
                _shade_row(row, "F5F5F5")

    # ═══════════════════════════════════════════════
    # Signature Block
    # ═══════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    for label in ["PREPARED BY:", "REVIEWED BY:"]:
        sig_label = doc.add_paragraph()
        run = sig_label.add_run(label)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Arial"
        sig_label.space_after = Pt(2)

        sig_fields = doc.add_paragraph()
        sig_fields.space_after = Pt(20)
        run = sig_fields.add_run("Name: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        sig_fields.add_run("_" * 30 + "   ")
        run = sig_fields.add_run("Signature: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        sig_fields.add_run("_" * 30 + "   ")
        run = sig_fields.add_run("Date: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        sig_fields.add_run("_" * 20)

    # Disclaimer
    doc.add_paragraph()
    disc = doc.add_paragraph()
    run = disc.add_run(
        "This report was auto-generated by the Pipe-Up Regulatory Compliance Engine. "
        "All crew locations and zone intersections are derived from daily work plan data "
        "and regulatory zone definitions. This report should be reviewed and verified by "
        "the responsible compliance officer before distribution."
    )
    run.font.size = Pt(8)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)
    run.italic = True

    doc.save(str(output_path))


def _add_section_heading(doc, text):
    """Add a navy-colored section heading (Heading 2)."""
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x1D, 0x33)
        run.font.name = "Arial"
    h.space_after = Pt(8)
    return h


def _apply_table_borders(table):
    """Apply thin borders to all cells in a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'</w:tblBorders>'
    )
    tbl_pr.append(borders)


def _style_table_header(row, labels):
    """Style a table header row with navy background and white text."""
    for i, label in enumerate(labels):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Arial"
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0B1D33"/>')
        cell._tc.get_or_add_tcPr().append(shading)


def _shade_cell(cell, hex_color):
    """Apply a background color to a single cell with white text."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True


def _shade_row(row, hex_color):
    """Apply alternating row shading (skip cells that already have shading)."""
    for cell in row.cells:
        # Check if cell already has custom shading (status/severity cells)
        tc_pr = cell._tc.tcPr
        if tc_pr is not None:
            existing = tc_pr.findall(qn('w:shd'))
            if existing:
                continue
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)


def _set_row_font(row, size):
    """Set font for all cells in a row."""
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = size
                run.font.name = "Arial"


# ─── Step 11: Main / Terminal Output ─────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Pipe-Up Daily Compliance Report Generator")
    parser.add_argument("--date", help="Process PDF for specific date (YYYY-MM-DD)")
    args = parser.parse_args()

    target_date = None
    if args.date:
        target_date = datetime.strptime(args.date, "%Y-%m-%d").date()

    # Step 1: Find PDF
    pdf_path, report_date = find_pdf(target_date)
    date_str = report_date.strftime("%B %d, %Y")
    date_file = report_date.strftime("%Y-%m-%d")

    # Extract pipeline name (strip "SMJV — " prefix if present)
    pipeline_name = CONFIG["project_name"]
    if "\u2014" in pipeline_name:
        pipeline_name = pipeline_name.split("\u2014", 1)[1].strip()
    elif "—" in pipeline_name:
        pipeline_name = pipeline_name.split("—", 1)[1].strip()

    BAR = "\u2550" * 43
    print(BAR)
    print("  Pipe-Up Daily Compliance Report")
    print(f"  {pipeline_name}")
    print(f"  {date_str}")
    print(BAR)

    # Step 2: Parse KML
    kp_points, centerline = parse_kml()
    interpolate_kp = build_kp_interpolator(kp_points)

    # Step 3 & 4: Parse PDF
    crews, crews_no_kp = parse_pdf(pdf_path)
    total_crews = len(crews) + len(crews_no_kp)
    total_kps = sum(len(c["kp_locations"]) for c in crews)
    print(f"  PDF Parsed:    {pdf_path.name}")
    print(f"  Crews Found:   {total_crews}")
    print(f"  KP Locations:  {total_kps}")

    # Add coordinates to crews
    for crew in crews:
        crew["coords"] = [interpolate_kp(kp) for kp in crew["kp_locations"]]

    # Step 7: Load zones and cross-reference
    zones, zone_config = load_zones()
    findings = cross_reference_zones(crews, zones, interpolate_kp)

    action_count = sum(1 for f in findings if f["severity"] == "HIGH")
    monitor_count = sum(1 for f in findings if f["severity"] == "MONITOR")
    compliant_count = sum(1 for f in findings if f["severity"] == "COMPLIANT")
    print(f"  Zone Crossings: {len(findings)}")
    print(f"    Action Required: {action_count}")
    print(f"    Monitor:         {monitor_count}")
    print(f"    Compliant:       {compliant_count}")

    # Step 8: Generate alerts
    alerts = generate_alerts(crews, zones, report_date, interpolate_kp)
    print(f"  Alerts:        {len(alerts)}")

    # Ensure output directory exists
    output_dir = BASE_DIR / CONFIG["output_dir"]
    output_dir.mkdir(exist_ok=True)

    # Step 9: Generate HTML map
    map_path = output_dir / f"EGP_Daily_Map_{date_file}.html"
    generate_html_map(crews, zones, zone_config, findings, alerts,
                      centerline, kp_points, interpolate_kp,
                      report_date, pdf_path.name, map_path)

    # Step 10: Generate Word report
    docx_path = output_dir / f"EGP_Compliance_Report_{date_file}.docx"
    generate_word_report(crews, zones, zone_config, findings, alerts,
                         report_date, pdf_path.name, docx_path)

    # Relative output paths for cleaner display
    map_rel = map_path.relative_to(BASE_DIR)
    docx_rel = docx_path.relative_to(BASE_DIR)

    print(BAR)
    print("  Output:")
    print(f"    \u2192 {map_rel}")
    print(f"    \u2192 {docx_rel}")
    print(BAR)


if __name__ == "__main__":
    main()
